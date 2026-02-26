"""
Sophisticated hierarchical recursive chunking strategy.

Pipeline:
  1. Load document with structure-aware loaders.
  2. First-pass split: break on semantic boundaries (double-newlines, headings).
  3. Second-pass split: RecursiveCharacterTextSplitter guarantees max chunk size.
  4. Merge orphan micro-chunks (<200 chars) into their predecessor.
  5. Enrich every chunk with rich metadata.
"""

import hashlib
import os
import re
from pathlib import Path
from typing import List

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from config import settings


# ── Loaders ──────────────────────────────────────────────────────────────────

def _load_pdf(path: str) -> List[Document]:
    import pdfplumber
    docs: List[Document] = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                docs.append(Document(
                    page_content=text,
                    metadata={"source": Path(path).name, "page": i}
                ))
    return docs


def _load_docx(path: str) -> List[Document]:
    from docx import Document as DocxDocument
    doc = DocxDocument(path)
    sections: List[str] = []
    current: List[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            if current:
                sections.append("\n".join(current))
                current = []
        else:
            current.append(text)
    if current:
        sections.append("\n".join(current))

    return [
        Document(
            page_content=s,
            metadata={"source": Path(path).name, "page": i + 1}
        )
        for i, s in enumerate(sections) if s.strip()
    ]


def _load_txt(path: str) -> List[Document]:
    text = Path(path).read_text(encoding="utf-8", errors="ignore")
    paragraphs = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]
    return [
        Document(
            page_content=p,
            metadata={"source": Path(path).name, "page": i + 1}
        )
        for i, p in enumerate(paragraphs)
    ]


def _load_md(path: str) -> List[Document]:
    text = Path(path).read_text(encoding="utf-8", errors="ignore")
    # Split on headings and blank-line gaps
    blocks = re.split(r"(?=^#{1,6} |\n{2,})", text, flags=re.MULTILINE)
    docs: List[Document] = []
    for i, block in enumerate(blocks):
        block = block.strip()
        if block:
            docs.append(Document(
                page_content=block,
                metadata={"source": Path(path).name, "page": i + 1}
            ))
    return docs


LOADERS = {
    ".pdf":  _load_pdf,
    ".docx": _load_docx,
    ".txt":  _load_txt,
    ".md":   _load_md,
}


# ── Chunker ───────────────────────────────────────────────────────────────────

_splitter = RecursiveCharacterTextSplitter(
    chunk_size=settings.CHUNK_SIZE,
    chunk_overlap=settings.CHUNK_OVERLAP,
    separators=["\n\n", "\n", ". ", "! ", "? ", ", ", " ", ""],
    length_function=len,
    add_start_index=True,
)


def _merge_micro_chunks(chunks: List[Document], min_len: int = 200) -> List[Document]:
    """Merge tiny trailing chunks into their predecessor."""
    merged: List[Document] = []
    for chunk in chunks:
        if merged and len(chunk.page_content) < min_len:
            merged[-1].page_content += " " + chunk.page_content
        else:
            merged.append(chunk)
    return merged


def load_and_chunk(file_path: str) -> List[Document]:
    """
    Load a document and return enriched chunks using the sophisticated
    hierarchical recursive chunking strategy.
    """
    ext = Path(file_path).suffix.lower()
    loader = LOADERS.get(ext)
    if loader is None:
        raise ValueError(f"Unsupported file type: {ext}")

    # Stage 1 – structure-aware load (preserves page / section boundaries)
    raw_docs = loader(file_path)

    # Stage 2 – recursive character splitting within each structural unit
    split_docs = _splitter.split_documents(raw_docs)

    # Stage 3 – merge micro-chunks
    final_chunks = _merge_micro_chunks(split_docs)

    # Stage 4 – enrich metadata
    source_name = Path(file_path).name
    total = len(final_chunks)
    for idx, doc in enumerate(final_chunks):
        doc.metadata.update({
            "source": source_name,
            "chunk_index": idx,
            "total_chunks": total,
        })

    return final_chunks


def compute_content_hash(file_path: str) -> str:
    """
    Compute a SHA-256 fingerprint of the document's full extracted text.

    Uses normalised text (not raw bytes) so the same content uploaded
    under a different filename — or with minor encoding differences —
    still produces the same hash.
    """
    ext = Path(file_path).suffix.lower()
    loader = LOADERS.get(ext)
    if loader is None:
        return ""
    docs = loader(file_path)
    # Normalise: join all text, collapse whitespace
    full_text = " ".join(" ".join(d.page_content.split()) for d in docs)
    return hashlib.sha256(full_text.encode("utf-8")).hexdigest()


def extract_preview_text(file_path: str, sample_chars: int = 1000) -> str:
    """
    Extract text samples from the START, MIDDLE, and END of the document
    so domain validation covers the full content, not just the opening lines.

    Each section contributes up to `sample_chars` characters.
    Returns a labelled string ready for the domain validator prompt.
    """
    ext = Path(file_path).suffix.lower()
    loader = LOADERS.get(ext)
    if loader is None:
        return ""

    docs = loader(file_path)
    if not docs:
        return ""

    # Flatten all pages/sections into one string
    full_text = "\n\n".join(d.page_content for d in docs).strip()
    total = len(full_text)

    if total == 0:
        return ""

    # ── Three non-overlapping windows ────────────────────────────────────────
    # START  : first sample_chars characters
    start_text = full_text[:sample_chars].strip()

    # MIDDLE : centred around the midpoint
    mid = total // 2
    mid_start = max(0, mid - sample_chars // 2)
    mid_text = full_text[mid_start: mid_start + sample_chars].strip()

    # END    : last sample_chars characters
    end_text = full_text[max(0, total - sample_chars):].strip()

    return (
        f"[START OF DOCUMENT]\n{start_text}\n\n"
        f"[MIDDLE OF DOCUMENT]\n{mid_text}\n\n"
        f"[END OF DOCUMENT]\n{end_text}"
    )
